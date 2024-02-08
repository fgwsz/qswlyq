# -*- coding: utf-8 -*-
import os
from PIL import Image
from PIL import ImageGrab
import docx
from datetime import datetime

def get_root_path():
    return os.path.dirname(os.path.abspath(__file__))

def get_current_date():
    date=datetime(datetime.now().year,datetime.now().month,datetime.now().day,0,0,0)
    return date.strftime("%Y%m%d")

def get_file_path():
    return get_root_path()+"\\"+get_current_date()+".docx"

def count_non_empty_cells_in_column(table,column_index):
    count=0
    for row in table.rows:
        cell=row.cells[column_index]
        if cell.text.strip():
            count+=1
    return count

if __name__=="__main__":
    image_path=get_root_path()+"\\__clipboard_image__.png"
    image=ImageGrab.grabclipboard()
    if image is None:
        print("剪切板中没有图像数据")
        exit()
    else:
        # 保存图像到当前路径下
        image.save(image_path,"PNG")
    document=docx.Document(get_file_path())
    table=document.tables[0]
    cell_col=1
    cell_row=count_non_empty_cells_in_column(table,cell_col)
    if(cell_row==0):
        exit()
    cell=table.cell(cell_row-1,cell_col)
    paragraph=cell.add_paragraph()
    paragraph.add_run().add_picture(image_path,width=cell.width)
    document.save(get_file_path())
    if os.path.exists(image_path):
        os.remove(image_path)
    print(f"{get_file_path()} table[0] index[{cell_row-1}] 添加图片成功!")
