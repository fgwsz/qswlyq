# -*- coding: utf-8 -*-
import os
import re
from datetime import datetime
import docx
import pyperclip

def get_root_path():
    return os.path.dirname(os.path.abspath(__file__))

def get_current_date():
    date=datetime(datetime.now().year,datetime.now().month,datetime.now().day,0,0,0)
    return date.strftime("%Y%m%d")

def get_file_path():
    return get_root_path()+"\\"+get_current_date()+".docx"

def count_non_empty_cells_in_column(table,column_index):
    count=0
    for row in table.rows:
        cell=row.cells[column_index]
        if cell.text.strip():
            count+=1
    return count

if __name__=="__main__":
    document=docx.Document(get_file_path())
    table_1=document.tables[0]
    table_2=document.tables[1]
    text=""
    text=text+f"{get_current_date()}报送情况统计\n"
    text=text+f"市级报送信息数量:{count_non_empty_cells_in_column(table_1,1)}\n"
    text=text+f"省级报送信息数量:{count_non_empty_cells_in_column(table_2,1)}\n"
    print(text)
    pyperclip.copy(text)
