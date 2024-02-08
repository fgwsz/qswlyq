# -*- coding: utf-8 -*-
import os
import re
from datetime import datetime
import docx
import pyperclip

def get_root_path():
    return os.path.dirname(os.path.abspath(__file__))

# 控制台读入一个指定格式的日期
# 如果输入为空，自动设置日期为今天的凌晨0点
def input_date(data_format,info):
    print(info)
    while True:
        date_str=input("请输入日期（格式为yyyyMMdd）：")
        if date_str.strip()=="":
            date=datetime(datetime.now().year,datetime.now().month,datetime.now().day,0,0,0)
            print("读入的日期是：",date)
            return date.strftime("%Y%m%d")
        else:
            try:
                date=datetime.strptime(date_str,date_format)
                print("读入的日期是：",date)
                return date_str
            except ValueError:
                print("日期格式不正确，请重新输入！")
                continue
# 得到当前文件夹下文件名在日期区间[start_date,end_date]之内的文件名列表
def get_file_list(start_date,end_date):
    start_date=datetime.strptime(start_date,'%Y%m%d')
    end_date=datetime.strptime(end_date,'%Y%m%d')
    file_list=[]
    for file in os.listdir(get_root_path()):
        if file.endswith('.docx')and re.match(r'\d{8}',file):
            file_date=datetime.strptime(file[:8],'%Y%m%d')
            if start_date<=file_date<=end_date:
                file_list.append(file)
    return file_list
# 统计docx.Document.tables[xxx]中下标为column_index的列中非空单元格的数量
def count_non_empty_cells_in_column(table,column_index):
    count=0
    for row in table.rows:
        cell=row.cells[column_index]
        if cell.text.strip():
            count+=1
    return count
# 统计信息(业务逻辑)
# docx文件中存在的表格信息:
# 第1个(index=0)表格的第2列(index=1)表示市报送信息数量
# 第2个(index=0)表格的第2列(index=1)表示省报送信息数量
def count_docx_info(file_path):
    ret=[]
    if file.endswith('.docx'):
        docx_file=docx.Document(file_path)
        tables=docx_file.tables
        row=0
        count=0
        print("文件:",file_path)
        for table in tables:
            count=count_non_empty_cells_in_column(table,1)
            if row==0:
                print("市级报送信息数量:",count)
            if row==1:
                print("省级报送信息数量:",count)
            ret.append(count)
            row+=1
    return ret

if __name__=="__main__":
    date_format="%Y%m%d"
    begin_date=input_date(date_format,"请输入开始日期")
    finish_date=input_date(date_format,"请输入截止日期")
    files=get_file_list(begin_date,finish_date)
    count=[]
    count.append(0)
    count.append(0)
    for file in files:
        file_count=count_docx_info(get_root_path()+"\\"+file)
        if len(file_count)==2:
            count[0]+=file_count[0]
            count[1]+=file_count[1]
    print("区间日期[",begin_date,",",finish_date,"]内报送情况统计")
    print("市级报送信息数量总计:",count[0])
    print("省级报送信息数量总计:",count[1])
