# -*- coding: utf-8 -*-
import os
from datetime import datetime
import docx

def get_root_path():
    return os.path.dirname(os.path.abspath(__file__))

def get_current_date():
    date=datetime(datetime.now().year,datetime.now().month,datetime.now().day,0,0,0)
    return date.strftime("%Y%m%d")

def get_file_path():
    return get_root_path()+"\\"+get_current_date()+".docx"

def count_non_empty_cells_in_column(table,column_index):
    count=0
    for row in table.rows:
        cell=row.cells[column_index]
        if cell.text.strip():
            count+=1
    return count

if __name__=="__main__":
    document=docx.Document(get_file_path())
    table=document.tables[0]
    cell_col=1
    cell_row=count_non_empty_cells_in_column(table,cell_col)
    if(cell_row==0):
        exit()
    cell=table.cell(cell_row-1,cell_col)
    # 清空单元格中的所有内容
    for paragraph in cell.paragraphs:
        # 清空段落中的文本
        paragraph.clear()
        # 删除段落中的图片
        for run in paragraph.runs:
            for picture in run.inline_shapes:
                paragraph._p.remove(run._r)
    # 获取单元格中的所有段落
    paragraphs=cell.paragraphs
    # 清空单元格中的所有段落，只保留一个段落
    while len(paragraphs)>1:
        cell._element.remove(paragraphs[-1]._element)
        del paragraphs[-1]
    document.save(get_file_path())
    print(f"{get_file_path()} table[0] index[{cell_row-1}] 清空成功!")
